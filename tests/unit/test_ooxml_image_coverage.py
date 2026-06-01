"""Complete embedded-image coverage for docx + xlsx.

docx: inline images (positioned) plus a package sweep that catches everything
the paragraph walk can't reach (anchored/floating, tables, headers/footers).
xlsx: images are harvested straight from the OPC package (read-only openpyxl
never loads drawings).
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XlImage
from PIL import Image as PILImage

from voitta_rag_enterprise.services.parsers.docx_parser import DocxParser
from voitta_rag_enterprise.services.parsers.xlsx_parser import XlsxParser


def _png(size: tuple[int, int] = (64, 64), color=(10, 120, 200)) -> bytes:
    buf = io.BytesIO()
    PILImage.new("RGB", size, color).save(buf, format="PNG")
    return buf.getvalue()


# --------------------------------------------------------------------------
# docx
# --------------------------------------------------------------------------

def test_docx_inline_image_extracted_and_positioned(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("before the picture")
    doc.add_picture(io.BytesIO(_png()))
    out = tmp_path / "inline.docx"
    doc.save(out)

    r = DocxParser().parse(out)
    assert r.success
    assert len(r.images) == 1
    assert r.images[0].mime in ("image/png",)
    assert (r.images[0].width, r.images[0].height) == (64, 64)


def test_docx_inline_image_not_double_counted(tmp_path: Path) -> None:
    """The package sweep must skip media already emitted by the inline walk."""
    doc = Document()
    doc.add_picture(io.BytesIO(_png()))
    out = tmp_path / "one.docx"
    doc.save(out)

    r = DocxParser().parse(out)
    assert len(r.images) == 1  # not 2


def test_docx_sweep_catches_non_inline_media(tmp_path: Path) -> None:
    """An image present in word/media/ but unreachable by the paragraph walk
    (anchored/floating, table, header/footer) is still extracted — appended at
    position 0. Simulated by injecting a media part the body never references."""
    doc = Document()
    doc.add_paragraph("text only, no inline image")
    out = tmp_path / "anchored.docx"
    doc.save(out)
    with zipfile.ZipFile(out, "a") as z:
        z.writestr("word/media/floating1.png", _png(size=(80, 80)))

    r = DocxParser().parse(out)
    assert len(r.images) == 1
    assert r.images[0].position == 0
    assert (r.images[0].width, r.images[0].height) == (80, 80)


def test_docx_sweep_filters_tiny_glyphs(tmp_path: Path) -> None:
    """Decorative <32px media (bullets, rules) are dropped by the sweep."""
    doc = Document()
    doc.add_paragraph("text")
    out = tmp_path / "tiny.docx"
    doc.save(out)
    with zipfile.ZipFile(out, "a") as z:
        z.writestr("word/media/bullet.png", _png(size=(12, 12)))
        z.writestr("word/media/real.png", _png(size=(64, 64)))

    r = DocxParser().parse(out)
    assert len(r.images) == 1  # only the 64px one survives
    assert (r.images[0].width, r.images[0].height) == (64, 64)


# --------------------------------------------------------------------------
# xlsx
# --------------------------------------------------------------------------

def test_xlsx_embedded_image_extracted(tmp_path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sales"
    ws.append(["Region", "Revenue"])
    ws.append(["EMEA", 100])
    ws.add_image(XlImage(io.BytesIO(_png())), "D2")
    out = tmp_path / "with_image.xlsx"
    wb.save(out)

    r = XlsxParser().parse(out)
    assert r.success
    assert "## Sheet: Sales" in r.content  # text pass still works
    assert len(r.images) == 1
    assert r.images[0].position == 0
    assert (r.images[0].width, r.images[0].height) == (64, 64)


def test_xlsx_without_images_has_none(tmp_path: Path) -> None:
    wb = Workbook()
    wb.active.append(["a", "b"])
    out = tmp_path / "plain.xlsx"
    wb.save(out)

    r = XlsxParser().parse(out)
    assert r.success
    assert r.images == []
